import os
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document
import streamlit as st
from g4f.client import Client
import g4f
import tempfile

modelo = 'gpt-3.5-turbo'




def chatear(modelo, mensaje):
    client = Client()
    with st.spinner('Contactando con la IA ...'):
        response = client.chat.completions.create(
            model=modelo,
            messages=[
                {"role": "user", "content": mensaje+" . Recuerda responder en español directamente, sin cinluir frase de apertura o de cierre de tu respuesta."}
            ], 
        )
        respuesta = response.choices[0].message.content
    return respuesta

def listar_archivos(directorio):
    return [archivo for archivo in os.listdir(directorio) if archivo.endswith('.docx') or archivo.endswith('.pdf') or archivo.endswith('.txt')]

def leer_docx(archivo):
    document = Document(archivo)
    return "\n".join(paragraph.text for paragraph in document.paragraphs)

def leer_pdf(archivo):
    with open(archivo, 'rb') as f:
        reader = PdfReader(f)
        texto = ''.join(page.extract_text() for page in reader.pages)
        indice_inicio = texto.find("INFORME DE AUDITORIA")
        return texto[indice_inicio:] if indice_inicio != -1 else texto

def guardar_word(respuesta, titulo):
    st.write("Presiona el siguiente botón para guardar el documento en formato Word:")
    if st.button("Guardar Word"):
        document = Document()
        document.add_heading(titulo, level=1)
        document.add_paragraph(str(respuesta).replace(". ", ".\n "))
        document.save(titulo + '.docx')
        st.success(f"El documento '{titulo}.docx' ha sido guardado exitosamente.")

def extraer_hallazgos(texto):
    indice_inicio = texto.find("INFORME DE AUDITORIA")
    texto_recortado = texto[indice_inicio:]
    inicio_hallazgos = texto_recortado.find("4. HALLAZGOS")
    fin_hallazgos = texto_recortado.find("5. ANÁLISIS DE LA VISTA")
    return texto_recortado[inicio_hallazgos:fin_hallazgos].strip()

def analizar_hallazgos(hallazgos):
    capitulos = hallazgos.split("\n4.")
    df_hallazgo = {}
    for i, capitulo in enumerate(capitulos, start=1):
        mensaje = (f"Te voy a pasar hallazgos de auditoría y un requerimiento que arranca con la palabra INSTRUCCION. "
                   f"{capitulo}INSTRUCCION. Necesito analizar la solidez técnica de este informe de auditoría de TI. "
                   f"Los hallazgos de auditoría se estructuran de la siguiente manera: a) Numeración (que siempre comienza con 4.) y título, "
                   f"b) descripción, c) criterio o referencia a normativa (o buenas prácticas de TI) y d) impacto (que podría pasar de no hacerlo). "
                   f"Necesito que revises los hallazgos para evaluar si son sólidos técnicamente (Si el criterio contra el que se contrasta es el más adecuado) "
                   f"o si consideras que debería citarse otra normativa o buena práctica, citando norma, edición, sección y título, por ejemplo: "
                   f"'COBIT 4.1 edición 2007: ME4 Proporcionar Gobierno de TI', o si hay algo incorrecto o técnicamente poco sólido. "
                   f"La respuesta debe identificar el número de hallazgo (por ej. 4.1.1) y comenzar por eso: 4.1.1: tu respuesta. "
                   f"Luego incluir el análisis de los criterios y el impacto, recordando si se mencionan criterios o normas, el incluir la norma, la edición y la sección específica.")
        st.write(f'Analizando Hallazgo {i} ...............')
        respuesta = chatear(modelo, mensaje)
        st.write(f'Hallazgo {i}: {respuesta}')
        df_hallazgo[f'Capítulo {i}'] = respuesta

    df_hallazgo = pd.DataFrame(list(df_hallazgo.items()), columns=['capitulo', 'texto'])
    document = Document()
    document.add_heading('Análisis de Hallazgos', 0)
    for i in df_hallazgo.index:
        document.add_paragraph(df_hallazgo['capitulo'][i])
        document.add_heading('Texto:', level=1)
        document.add_paragraph(df_hallazgo['texto'][i])
    document.save('hallazgos.docx')
    st.write('Guardando Word hallazgos.docx')

def leer_archivo(tipo_archivo, archivo):
    if archivo is None:
        st.warning("Por favor, selecciona un archivo primero.")
        return
    
    # Obtener el nombre del archivo
    nombre_archivo = archivo.name
    
    # Guardar el archivo temporalmente
    with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
        tmp_file.write(archivo.read())
        tmp_file.seek(0)
        ruta_archivo_temporal = tmp_file.name
    
    if nombre_archivo.endswith('.docx'):
        texto = leer_docx(ruta_archivo_temporal)
    else:
        texto = leer_pdf(ruta_archivo_temporal)

    if tipo_archivo == "minutas":
        st.write("Buscando hallazgos en una minuta ...")
        mensaje = (f"Debes analizar esta minuta y buscar declaraciones o frases que pudieran significar directa o indirectamente un riesgo, amenaza o simplemente que no esté alineado a las buenas prácticas. "
                   f"La palabra confidencial en la minuta, significa borrador. Acto seguido deberás enumerarlas, mencionando la frase exacta y a continuación, relacionarlo con algún criterio o buena práctica como la ISO 27001, COBIT, ITIL, etc. "
                   f"Para cada uno deberás ser muy preciso en cuanto a la norma o buena práctica, la versión, y la sección específica como para que pueda fácilmente buscarla en internet y averiguar más al respecto. "
                   f"La minuta es la siguiente: {texto}")
        respuesta = chatear(modelo, mensaje)
        st.write(respuesta)
        guardar_word(respuesta, 'minuta_hallazgos')
    elif tipo_archivo == "informe":
        opcion = st.selectbox("Elige una opción:", ["Armar el glosario", "Validación ortográfica", "Validación técnica de hallazgos", "Escribir las conclusiones del informe", "Relacionar informe con los ODS"])
        if opcion == "Armar el glosario":
            st.write("Armando el Glosario...")
            mensaje = (f"Debes analizar este texto. El mismo es un informe de auditoría y debe poder ser leído por cualquier ciudadano. "
                       f"Busca palabras que creas que deberían estar en un glosario, en especial a las relacionadas con las tecnologías, protocolos, etc. "
                       f"Si el informe ya cuenta con un glosario, a ese glosario debes agregarle lo que consideres. "
                       f"Es muy importante que armes el glosario en orden alfabético. El texto es el siguiente: {texto}")
            respuesta = chatear(modelo, mensaje)
            st.write('Respuesta: ', respuesta)
            guardar_word(respuesta, 'Glosario')
        elif opcion == "Validación ortográfica":
            st.write("Validación ortográfica ...")
            mensaje = (f"Debes analizar este texto. El mismo es un informe de auditoría y NO debe contener errores de ortografía. "
                       f"Valida el texto según la RAE y si hay errores, sugiere cómo corregirlos. El texto es el siguiente: {texto}")
            respuesta = chatear(modelo, mensaje)
            st.write(respuesta)
            guardar_word(respuesta, 'validación_redaccion')
        elif opcion == "Validación técnica de hallazgos":
            st.write("Validación técnica de hallazgos ...")
            hallazgos = extraer_hallazgos(texto)
            analizar_hallazgos(hallazgos)
        elif opcion == "Escribir las conclusiones del informe":
            st.write("Escribiendo las conclusiones del informe...")
            hallazgos = extraer_hallazgos(texto)
            mensaje = (f"Debes analizar este informe. Debes resumir todo lo mencionado y redactar las conclusiones en un lenguaje llano como para que cualquiera lo entienda. "
                       f"El informe es el siguiente: {hallazgos}")
            respuesta = chatear(modelo, mensaje)
            guardar_word(respuesta, 'conclusiones')
        elif opcion == "Relacionar informe con los ODS":
            st.write("Buscando relaciones con los ODS ...")
            mensaje = (f"Debes analizar este informe y en base al OBJETO DE AUDITORIA, el organismo, su misión y función, y el alcance, debes encontrar relaciones con las metas específicas o indicadores de los ODS 2030 de la ONU. "
                       f"Tu informe debe incluir por ODS, la meta o indicador con su descripcion completa y una explicación detallada de cada relacion que encuentres. El informe es el siguiente: {texto}")
            respuesta = chatear(modelo, mensaje)
            st.write(respuesta)
            guardar_word(respuesta, 'ODS')




def main():
    st.image("header.jpg", use_column_width=True)
    st.write("Selecciona la opción:")
    opcion = st.selectbox("Elige una opción:", ["Trabajar con Minutas", "Trabajar con Informe Borrador", "Ayuda para Marco normativo", "Preguntas guía para entrevista de auditoría"])
    
    if opcion == "Trabajar con Minutas":
        uploaded_file = st.file_uploader("Elige un archivo", type=["docx", "pdf"])
        if uploaded_file is not None:
            leer_archivo("minutas", uploaded_file)
    elif opcion == "Trabajar con Informe Borrador":
        uploaded_file = st.file_uploader("Elige un archivo", type=["docx", "pdf"])
        if uploaded_file is not None:
            leer_archivo("informe", uploaded_file)
    elif opcion == "Ayuda para Marco normativo":
        organismo = st.text_input("Ingrese el Organismo:")
        objeto = st.text_input("Ingrese el objeto de la auditoría:")
        if st.button("Generar Marco Normativo"):
            mensaje = (f"Hola. Necesito ayuda para encontrar criterios o materia aplicable para confeccionar un marco normativo para que la AGN realice una auditoría de TI en {organismo} en Argentina. "
                       f"El objeto de auditoría es el {objeto}. Debes tener muy en cuenta las características del organismo y las leyes o normativas aplicables más útiles para auditar TI en Argentina. "
                       f"Al lado de cada sugerencia debes poner entre parentesis la relación en una escala de baja, media o alta. Quiero que seas lo más detallista posible al sugerir, por ejemplo si citas una norma, debes citar la norma exacta y ser lo más descriptivo posible. "
                       f"Es indispensable que me des la mayor cantidad de herramientas posibles. Solo lístame los elementos, no es necesario que armes el texto del marco normativo, ni que agregues textos de introducción o de cierre.")
            st.write("Generando referencias para el marco normativo ...")
            respuesta = chatear(modelo, mensaje)
            st.write('Respuesta: ', respuesta)
            guardar_word(respuesta, 'Marco_normativo')
    elif opcion == "Preguntas guía para entrevista de auditoría":
        organismo = st.text_input("Ingrese ORGANISMO auditado y el OBJETO de la auditoria:")
        entrevistado = st.text_input("Ingrese con quien será la entrevista:")
        if st.button("Generar Preguntas Guía"):
            mensaje = (f"Hola. Necesito ayuda para crear preguntas guía para realizar una entrevista en el marco de una auditoría. El organismo y objeto de auditoría son {organismo}. "
                       f"La entrevista sera con {entrevistado}, que trabaja en el organismo mencionado. Debes intuir por donde irá la reunión y crear todas las preguntas guía posibles para facilitarme el trabajo como auditor. Al final de cada pregunta, debes poner entre parentesis que relacion tiene (alto, medio, bajo) con el objeto de auditoria. Espero la lista de preguntas sin introducción ni cierre y que sean las más posibles.")
            st.write("Generando preguntas guía para la entrevista ...")
            respuesta = chatear(modelo, mensaje)
            st.write('Respuesta: ', respuesta)
            guardar_word(respuesta, 'Guia_entrevista')

if __name__ == "__main__":
    # Configuración de la página principal
    st.set_page_config(
        page_title="Mi Aplicación",
        layout="wide"
    )

    # Función para mostrar la página de inicio
    def home():
        st.image("home.png", use_column_width=True)
        st.title("Aviso Legal:")
        st.write("Bienvenido a nuestro asistente virtual asistido por inteligencia artificial para auditoría de TI. Antes de continuar, por favor lea detenidamente el siguiente aviso:")
        st.write("Propósito del Asistente: Este asistente virtual está diseñado para proporcionar ayuda y soporte en temas relacionados con la auditoría de TI. La información y recomendaciones proporcionadas son generadas por un sistema de inteligencia artificial y están destinadas a servir como guía y apoyo.")
        st.write("Revisión Profesional: La información proporcionada por este asistente no sustituye el juicio, análisis y revisión de un profesional calificado en auditoría de TI. Es esencial que cualquier información o recomendación obtenida a través de este asistente sea revisada y validada por un profesional antes de tomar decisiones o implementar cualquier acción basada en dicha información.")
        st.write("Limitación de Responsabilidad: No nos hacemos responsables por cualquier daño, pérdida o perjuicio que pueda derivarse del uso de la información proporcionada por este asistente virtual. El usuario asume toda la responsabilidad por el uso de la información y las decisiones tomadas con base en la misma.")
        st.write("Actualización y Precisión de la Información: Aunque nos esforzamos por proporcionar información precisa y actualizada, no garantizamos la exactitud, integridad o vigencia de la información proporcionada por el asistente. La tecnología y las mejores prácticas en auditoría de TI pueden cambiar, y la información puede quedar desactualizada.")
        st.write("Aceptación de los Términos: Al utilizar este asistente virtual, usted acepta y se compromete a cumplir con los términos y condiciones establecidos en este aviso legal y descargo de responsabilidad.")

    # Función para la página de Minutas
    def plani():
        st.image("plani.png", use_column_width=True)
        st.title("Planificación")
        st.write("Selecciona la opción:")
        opcion = st.selectbox("Elige una opción:", ["Asistencia para Marco normativo","Antecedentes relacionados","Preparar entrevista", "Analizar minutas de entrevista"])
        if opcion == "Analizar minutas de entrevista":
            uploaded_file = st.file_uploader("Elige un archivo", type=["docx", "pdf"])
            if uploaded_file is not None:
                leer_archivo("minutas", uploaded_file)
        elif opcion == "Asistencia para Marco normativo":
            organismo = st.text_input("Ingrese el Organismo:")
            objeto = st.text_input("Ingrese el objeto de la auditoría:")
            if st.button("Generar Marco Normativo"):
                mensaje = (f"Hola. Necesito ayuda para encontrar criterios o materia aplicable para confeccionar un marco normativo para que la AGN realice una auditoría de TI en {organismo} en Argentina. "
                        f"El objeto de auditoría es el {objeto}. Debes tener muy en cuenta las características del organismo y las leyes o normativas aplicables más útiles para auditar TI en Argentina. "
                        f"Al lado de cada sugerencia debes poner entre parentesis la relación en una escala de baja, media o alta. Quiero que seas lo más detallista posible al sugerir, por ejemplo si citas una norma, debes citar la norma exacta y ser lo más descriptivo posible. "
                        f"Es indispensable que me des la mayor cantidad de herramientas posibles. Solo lístame los elementos, no es necesario que armes el texto del marco normativo, ni que agregues textos de introducción o de cierre.")
                st.write("Generando referencias para el marco normativo ...")
                respuesta = chatear(modelo, mensaje)
                st.write('Respuesta: ', respuesta)
                guardar_word(respuesta, 'Marco_normativo')
        elif opcion == "Preparar entrevista":
            organismo = st.text_input("Ingrese ORGANISMO auditado y el OBJETO de la auditoria:")
            entrevistado = st.text_input("Ingrese con quien será la entrevista:")
            if st.button("Generar Preguntas Guía"):
                mensaje = (f"Hola. Necesito ayuda para crear preguntas guía para realizar una entrevista en el marco de una auditoría. El organismo y objeto de auditoría son {organismo}. "
                        f"La entrevista sera con {entrevistado}, que trabaja en el organismo mencionado. Debes intuir por donde irá la reunión y crear todas las preguntas guía posibles para facilitarme el trabajo como auditor. Al final de cada pregunta, debes poner entre parentesis que relacion tiene (alto, medio, bajo) con el objeto de auditoria. Espero la lista de preguntas sin introducción ni cierre y que sean las más posibles.")
                st.write("Generando preguntas guía para la entrevista ...")
                respuesta = chatear(modelo, mensaje)
                st.write('Respuesta: ', respuesta)
                guardar_word(respuesta, 'Guia_entrevista')
        elif opcion == "Antecedentes relacionados":
            organismo = st.text_input("Ingrese ORGANISMO auditado y el OBJETO de la auditoria:")
            if st.button("Verificar Antecedentes"):
                mensaje = (f"Hola. Eres auditor gubernamental, especialista en auditoría de TI. Quiero todos los antecedentes, informes previos, noticias periodísticas, etc que tengas para {organismo}."
                        f"Quiero que separes los antecedentes entre los relacionados con el organismo, y los relacionados con el objeto. Pueden ser auditorías previas de AGN, de otro órgano de control de Argentina, u otra EFS/SAI del mundo (debes especificar entre paréntesis la EFS). En cada relación debes especificar que relacion tiene (alto, medio, bajo) con el input dado. Espero tu trabajo sin que envíes nada de introducción ni cierre y que sea lo más extenso posible.")
                st.write("Verificando antecedentes ...")
                respuesta = chatear(modelo, mensaje)
                st.write('Respuesta: ', respuesta)
                guardar_word(respuesta, 'Antecedentes')
                
    # Función para la página de Informe
    def ejec():
        st.image("ejec.png", use_column_width=True)
        st.title("Analizar informe ...")
        uploaded_file = st.file_uploader("Elige un archivo", type=["docx", "pdf"])
        if uploaded_file is not None:
            leer_archivo("informe", uploaded_file)
        

    # Función para la página de Marco Normativo
    def informe():
        st.image("informe.png", use_column_width=True)
        st.title("Sobre el informe")
        st.write("Aquí puedes consultar el informe...")

    # Función para la página de Preguntas de Entrevista
    def otros():
        st.image("home.png", use_column_width=True)
        st.title("Otros")
        objeto = st.text_input("Ingrese la consulta:")
        if st.button("Enviar consulta"):
            mensaje = (f"Hola. actúa con un asistente profesional en auditoria de TI en la Auditoría General de la Nación. Debes responder esta consulta: {organismo}")
            st.write("Enviando consulta ...")
            respuesta = chatear(modelo, mensaje)
            st.write('Respuesta: ', respuesta)
            guardar_word(respuesta, 'Otras consultas')

    # Configuración del menú en la barra lateral
    menu_options = ["Inicio", "Planificación", "Ejecución", "Informe", "Otros"]
    selection = st.sidebar.selectbox("Menu", menu_options)

    # Navegación basada en la selección del menú
    if selection == "Inicio":
        home()
    elif selection == "Planificación":
        plani()
    elif selection == "Ejecución":
        ejec()
    elif selection == "Informe":
        informe()
    elif selection == "Otros":
        otros()