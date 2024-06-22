import streamlit as st
import PyPDF2
import docx
import g4f
from g4f.client import Client
from io import BytesIO


def extract_text_from_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_docx(file):
    doc = docx.Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def get_gpt_response(prompt):
    client = Client()
 
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt+". Recuerda responder en español directamente, sin inluir frase de apertura o de cierre de tu respuesta."}],
    )
    return response.choices[0].message.content

def get_book_summary(book_text, title):
    prompt = f"Resumen de no menos de 1000 palabras del libro '{title}' con pasajes destacados:\n\n{book_text[:2000]}"
    return get_gpt_response(prompt)

def get_work_tips(summary, profession):
    prompt = f"Basado en el siguiente resumen de libro:\n{summary}\n\nCómo puede agregar valor al trabajo de un {profession}:"
    return get_gpt_response(prompt)

def get_poner_en_practica(title):
    prompt = f"Que ejercicio o dinámica me recomiendas con mis subordinados, para poner en práctica lo aprendido en '{title}':"
    return get_gpt_response(prompt)

def get_similar_books(title):
    prompt = f"Libros similares o complementarios a '{title}':"
    return get_gpt_response(prompt)

def create_word_document(content):
    doc = docx.Document()
    doc.add_paragraph(content)
    return doc

st.title("Analizador de Libros con IA")

book_title = st.text_input("Título del libro:")
uploaded_file = st.file_uploader("O sube el archivo del libro (PDF o DOCX):", type=["pdf", "docx"])
profession = st.text_input("Tu profesión o puesto:",value="Operations and Finance Manager - Natura&Co Pay en Natura &Co")

if st.button("Analizar"):
    if uploaded_file:
        if uploaded_file.type == "application/pdf":
            book_text = extract_text_from_pdf(uploaded_file)
        else:
            book_text = extract_text_from_docx(uploaded_file)
        title = uploaded_file.name
    elif book_title:
        book_text = f"Contenido del libro '{book_title}'"
        title = book_title
    else:
        st.error("Por favor, ingresa un título o sube un archivo.")
        st.stop()

    with st.spinner("Analizando..."):
        summary = get_book_summary(book_text, title)
        tips = get_work_tips(summary, profession)
        practica = get_poner_en_practica(title)
        similar_books = get_similar_books(title)

        st.subheader("Resumen del libro y pasajes destacados:")
        st.write(summary)

        st.subheader("Cómo puede agregar valor a tu trabajo:")
        st.write(tips)
        st.subheader("Cómo ponerlo en práctica con mi equipo:")
        st.write(practica)

        st.subheader("Libros similares o complementarios:")
        st.write(similar_books)
        st.markdown("¿Quieres apoyarme con un café? ¡Haz clic en el siguiente enlace!")
        st.markdown('<a href="https://cafecito.app/marianoparada" target="_blank"><img src="https://cdn.cafecito.app/imgs/buttons/button_2.svg" alt="Invitame un café en cafecito.app" /></a>', unsafe_allow_html=True)
        # Crear documento de Word
        doc_content = f"Resumen de '{title}':\n\n{summary}\n\nCómo agregar valor al trabajo de un {profession}:\n\n{tips}\n\nLibros similares o complementarios:\n\n{similar_books}"
        doc = create_word_document(doc_content)

        # Botón de descarga
        bio = BytesIO()
        doc.save(bio)
        st.download_button(
            label="Descargar resumen en Word",
            data=bio.getvalue(),
            file_name=title+"_resumen_libro.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        
